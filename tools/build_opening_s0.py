from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(r"D:\project\world-loli")
OUT = ROOT / "doc" / "剧情设计" / "开局" / "斗萝大陆_分场脚本_S0_冷开场_月痕_V1.0.docx"

INK, BLUE, MUTED = "172033", "2E5E78", "66737F"
PALE, GRAY, WHITE, RED = "E8F0F4", "F3F5F7", "FFFFFF", "8E2F3C"


def font(run, size=10.5, bold=False, color=INK, italic=False):
    name = "Microsoft YaHei"
    run.font.name = name
    run._element.get_or_add_rPr()
    for key in ("ascii", "hAnsi", "eastAsia"):
        run._element.rPr.rFonts.set(qn(f"w:{key}"), name)
    run.font.size = Pt(size)
    run.bold, run.italic = bold, italic
    run.font.color.rgb = RGBColor.from_string(color)


def para(p, before=0, after=6, line=1.22, align=None):
    f = p.paragraph_format
    f.space_before, f.space_after, f.line_spacing = Pt(before), Pt(after), line
    if align is not None:
        p.alignment = align


def text(doc, value, size=10.5, bold=False, color=INK, italic=False,
         before=0, after=6, align=None):
    p = doc.add_paragraph()
    para(p, before, after, 1.22, align)
    font(p.add_run(value), size, bold, color, italic)
    return p


def labeled(doc, label, value, after=5):
    p = doc.add_paragraph()
    para(p, 0, after, 1.22)
    font(p.add_run(label), 10.5, True, BLUE)
    font(p.add_run(value), 10.5)


def heading(doc, value, level=1):
    p = doc.add_paragraph()
    para(p, 14 if level == 1 else 9, 7 if level == 1 else 4, 1.1)
    font(p.add_run(value), 15 if level == 1 else 12, True, BLUE)
    return p


def shade(cell, fill):
    pr = cell._tc.get_or_add_tcPr()
    shd = pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        pr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell, top=85, start=110, bottom=85, end=110):
    pr = cell._tc.get_or_add_tcPr()
    node = pr.first_child_found_in("w:tcMar")
    if node is None:
        node = OxmlElement("w:tcMar")
        pr.append(node)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        item = node.find(qn(f"w:{name}"))
        if item is None:
            item = OxmlElement(f"w:{name}")
            node.append(item)
        item.set(qn("w:w"), str(value))
        item.set(qn("w:type"), "dxa")


def geometry(table, widths, indent=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    pr = table._tbl.tblPr
    width = pr.first_child_found_in("w:tblW")
    if width is None:
        width = OxmlElement("w:tblW")
        pr.append(width)
    width.set(qn("w:w"), str(sum(widths)))
    width.set(qn("w:type"), "dxa")
    ind = pr.first_child_found_in("w:tblInd")
    if ind is None:
        ind = OxmlElement("w:tblInd")
        pr.append(ind)
    ind.set(qn("w:w"), str(indent))
    ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for old in list(grid):
        grid.remove(old)
    for value in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(value))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tcw = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            if tcw is None:
                tcw = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tcw)
            tcw.set(qn("w:w"), str(widths[i]))
            tcw.set(qn("w:type"), "dxa")


def borders(table, color="CBD4DA", size="4"):
    pr = table._tbl.tblPr
    node = pr.first_child_found_in("w:tblBorders")
    if node is None:
        node = OxmlElement("w:tblBorders")
        pr.append(node)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        item = node.find(qn(f"w:{edge}"))
        if item is None:
            item = OxmlElement(f"w:{edge}")
            node.append(item)
        item.set(qn("w:val"), "single")
        item.set(qn("w:sz"), size)
        item.set(qn("w:space"), "0")
        item.set(qn("w:color"), color)


def keep_rows_together(table):
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))


def callout(doc, label, value, warning=False):
    table = doc.add_table(rows=1, cols=1)
    geometry(table, [9360])
    borders(table, "D3DEE5")
    cell = table.cell(0, 0)
    shade(cell, "F8ECEE" if warning else PALE)
    margins(cell, 130, 160, 130, 160)
    p = cell.paragraphs[0]
    para(p, 0, 0, 1.2)
    font(p.add_run(label + " "), 10, True, RED if warning else BLUE)
    font(p.add_run(value), 10)


def metadata(doc):
    rows = [
        ("场次", "S0", "场名", "冷开场：月痕"),
        ("场景", "微光森林深处·古树领地", "时间", "黄昏将尽"),
        ("时长", "约90秒（60–120秒弹性）", "触发", "首次创建角色后"),
        ("角色", "实体露娜", "交互", "纯演出，可跳过"),
    ]
    table = doc.add_table(rows=0, cols=4)
    geometry(table, [1100, 3460, 1100, 3700])
    borders(table)
    for vals in rows:
        cells = table.add_row().cells
        for i, value in enumerate(vals):
            cells[i].text = ""
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            margins(cells[i], 90, 110, 90, 110)
            if i in (0, 2):
                shade(cells[i], PALE)
            p = cells[i].paragraphs[0]
            para(p, 0, 0, 1.1)
            font(p.add_run(value), 9.3, i in (0, 2), BLUE if i in (0, 2) else INK)


def shot_table(doc):
    shots = [
        ("01", "0–5s", "黑场", "黑暗中先出现极淡雾粒，不展示环境。", "无", "风声、树叶摩擦；远处狼嚎被雾吞没。"),
        ("02", "5–13s", "航拍缓降", "树梢被风吹向东，林下冷雾却贴地向西倒卷，集中流向古树。", "无", "风与雾形成相反声向；低频倒吸声进入。"),
        ("03", "13–20s", "低机位近景", "露娜的靴子踏入画面，落地踉跄半步。一滴带银光的血落在苔藓上，苔藓微光熄灭。", "无", "枯枝断裂；压低的喘息。"),
        ("04", "20–30s", "中近景跟拍", "露娜扶住左肩穿过林间。银发与狼耳被逆流雾拂向错误方向。她回望身后，确认追踪者仍未现身。", "无", "衣料摩擦；狼嚎变远。音乐只给不完整单音。"),
        ("05", "30–40s", "主观扫视", "无风自颤的叶片、突然中断的狼族足迹、汇入古树根部的雾线依次入焦。", "无", "环境声逐层抽离，只剩两拍心跳。"),
        ("06", "40–52s", "侧面中景", "露娜走到古树与狼族领地之间，强迫自己站直；触地辨认气味后迅速收手。", "露娜（极轻）：\n“不是野兽的味道……”", "台词后留一秒静默。"),
        ("07", "52–62s", "面部特写", "露娜狼耳转向林中异响。她护住伤肩但不后退，视线追向雾的来处。", "露娜（压住喘息）：\n“断月纹还在追我。”", "林深处传来似脚步又似树根断裂的闷响。"),
        ("08", "62–74s", "树根特写", "雾从树根缝隙被吸入，三段冷白刻线依次亮起，拼成残缺月牙；缺口带明显人为切断感。", "露娜（画外）：\n“有人把不属于森林的东西埋进来了。”", "每亮一段伴随细小石裂音。"),
        ("09", "74–84s", "伤口极近景", "露娜伸手接近刻痕。左肩旧伤渗出断续银色雾光，被刻痕牵引成短线。", "无", "虫鸣瞬间停止；低频共振升高。"),
        ("10", "84–90s", "正面近景→黑场", "刻痕爆出冷白光。露娜抬臂护面，另一只手仍撑在古树前。白光吞没画面后硬切黑。", "字幕：\n“同一时刻，晨曦村东侧村道。”", "村钟响第一声，尾音延续进入S1。"),
    ]
    table = doc.add_table(rows=1, cols=6)
    geometry(table, [520, 680, 1100, 2700, 2220, 2140])
    borders(table)
    headers = ["镜号", "时间", "景别/机位", "画面与表演", "对白/字幕", "声音/转场"]
    for i, value in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade(cell, BLUE)
        margins(cell, 90, 85, 90, 85)
        p = cell.paragraphs[0]
        para(p, 0, 0, 1.0, WD_ALIGN_PARAGRAPH.CENTER)
        font(p.add_run(value), 8.6, True, WHITE)
    for row_no, values in enumerate(shots, 1):
        cells = table.add_row().cells
        for i, value in enumerate(values):
            if row_no % 2 == 0:
                shade(cells[i], GRAY)
            margins(cells[i], 80, 85, 80, 85)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            p = cells[i].paragraphs[0]
            para(p, 0, 0, 1.08)
            font(p.add_run(value), 8.1, i == 0, BLUE if i == 0 else INK)
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    keep_rows_together(table)


doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.top_margin = sec.bottom_margin = Inches(0.72)
sec.left_margin = sec.right_margin = Inches(1)
sec.header_distance, sec.footer_distance = Inches(0.35), Inches(0.35)

normal = doc.styles["Normal"]
normal.font.name, normal.font.size = "Microsoft YaHei", Pt(10.5)
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
normal.paragraph_format.space_after, normal.paragraph_format.line_spacing = Pt(6), 1.25

footer = sec.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
font(footer.add_run("《斗萝大陆》分场脚本 · S0 冷开场：月痕 · V1.0"), 8, False, MUTED)

text(doc, "《斗萝大陆》分场脚本", 22, True, INK, after=3)
text(doc, "S0  冷开场：月痕", 17, True, BLUE, after=12)
metadata(doc)
text(doc, "版本依据：《斗萝大陆_开局剧情设计_V2.0》；V3.1 仅参考逐镜头表达粒度。", 9, False, MUTED, True, 5, 9)

heading(doc, "一、场次目的")
labeled(doc, "剧情：", "用最少信息建立森林规则被改写、露娜已为此付出代价，并留下“断月纹从何而来”的长期悬念。")
labeled(doc, "角色：", "露娜首次登场即表现出敏锐、克制和守护责任。她受伤但不求救，始终把自己放在古树与狼族领地之间。")
labeled(doc, "转场：", "以冷白闪光切黑，再用晨曦村村钟把 S0 与 S1 压在同一时间线上。")
callout(doc, "信息边界", "不解释源萝之力、幕后污染者或月痕全貌；不出现任务、战斗、奖励或契约信息。跳过本场不会丢失关键操作信息。")

heading(doc, "二、逐镜头脚本")
shot_table(doc)

heading(doc, "三、对白与表演定稿")
text(doc, "露娜（辨认雾线，极轻）：不是野兽的味道……", after=3)
text(doc, "露娜（确认追踪仍在持续，压住喘息）：断月纹还在追我。", after=3)
text(doc, "露娜（看向古树根部，警惕压过愤怒）：有人把不属于森林的东西埋进来了。", after=7)
callout(doc, "表演原则", "露娜的核心状态是“负伤后仍在执行守护判断”。喘息只表现身体代价，不能把她演成等待救援的弱者。最后护挡动作必须面向狼族领地方向。")

heading(doc, "四、音画与程序执行")
for label, value in [
    ("风雾机制：", "树梢、落叶、露娜发丝向东；贴地雾粒向西并汇入古树。预算不足时也必须保住至少两组反向运动参照。"),
    ("断月纹：", "冷白、残缺、带人为切断缺口；不可画成完整弯月、狼族图腾或温暖金色法阵。"),
    ("月痕伤口：", "只显示断续银色雾光与牵引细线，不做大面积发光纹身或华丽技能爆发。"),
    ("跳过规则：", "开场3秒后允许跳过；落点为S1首帧，村钟尾音保留一次。重播不得触发任务、奖励或状态写入。"),
    ("字幕：", "对白全程带字幕；结尾仅显示“同一时刻，晨曦村东侧村道。”，不添加世界观说明旁白。"),
]:
    labeled(doc, label, value)

heading(doc, "五、贴图缺口与生成提示词")
text(doc, "现有可复用：露娜基础造型、古树单体、月光森林氛围背景。以下资产为 S0 制作缺口，提示词默认要求与现有二次元奇幻、清晰轮廓、青蓝森林色调一致。", 9.5, False, MUTED, True, after=7)

assets = [
    ("A01 必需", "微光森林深处·逆流雾古树背景", "16:9，建议 1920×1080，无遮挡背景",
     "二次元奇幻游戏剧情背景，微光森林深处，黄昏将尽，巨大古树位于画面偏右，树根盘错并留出角色站位，树梢与落叶明显向东倾斜，贴地银蓝冷雾却逆向向西流动并汇聚到古树根部，远处森林幽深，没有可见敌人，青蓝与深绿色主色，少量冷白微光，神秘压迫但不恐怖，精致游戏CG背景，清晰景深，适合横屏过场动画，无角色，无文字，无UI，无月亮法阵"),
    ("A02 必需", "负伤露娜·守护姿态演出立绘", "透明背景，竖版全身或大腿以上",
     "沿用参考角色露娜的固定设计：娇小银灰发狼耳少女，金色眼睛，侧马尾，深青蓝斗篷与月牙饰品；肩侧受伤，衣料有轻微破损和少量血迹，左手压住伤肩，身体略微失衡但强迫站直，右臂向外护挡，警惕看向画面外威胁，表情克制坚决而非痛哭求救，发丝和斗篷被反向雾流吹动，二次元奇幻RPG剧情立绘，清晰完整轮廓，透明背景，无文字，无UI，无额外角色，无性感化，无夸张伤口"),
    ("A03 必需", "残缺断月刻痕特效", "透明背景，正视图，可叠加",
     "透明背景游戏特效贴图，古老石刻般的残缺月牙符文，由三段不连续的冷白刻线组成，月牙缺口具有明显人为切断和侵蚀感，边缘散发少量银蓝雾粒与细小石屑，中心保持透明，不是完整月亮，不是魔法阵，不含文字，不含地面，不含角色，高清锐利边缘，适合加色与发光叠加"),
    ("A04 必需", "露娜肩伤·银雾月痕特效", "透明背景，局部特效",
     "透明背景局部游戏特效，肩部伤口逸散的断续银蓝雾光，细小雾丝从伤口向一个方向被牵引，带少量冷白闪点，表现侵蚀和共鸣而不是技能释放，克制、危险、低亮度，边缘柔和，中间透明，无人体、无衣服、无文字、无完整法阵、无火焰"),
    ("A05 可选", "逆流雾前景叠加层", "透明背景，16:9",
     "透明背景横向雾效叠加素材，银蓝色贴地薄雾从右向左逆流，包含前中后三层速度感和少量发光雾粒，上半部大面积透明，柔和边缘，可无缝循环，二次元奇幻游戏特效，无场景，无角色，无文字，无UI"),
    ("A06 可选", "白光吞没转场", "透明或黑底序列关键帧，16:9",
     "横屏游戏过场转场特效，冷白光从残缺月牙形裂隙瞬间扩散并吞没画面，边缘带极淡银蓝雾晕，中心高亮但保留层次，最后可过渡到纯黑，克制神秘，不要爆炸火焰，不要完整魔法阵，不要文字，不要角色，不要UI"),
]
table = doc.add_table(rows=1, cols=4)
geometry(table, [980, 1800, 1500, 5080])
borders(table)
for i, value in enumerate(("编号", "资产", "规格", "正向提示词")):
    cell = table.rows[0].cells[i]
    shade(cell, BLUE)
    margins(cell)
    p = cell.paragraphs[0]
    para(p, 0, 0, 1.0, WD_ALIGN_PARAGRAPH.CENTER)
    font(p.add_run(value), 8.7, True, WHITE)
for row_no, values in enumerate(assets, 1):
    cells = table.add_row().cells
    for i, value in enumerate(values):
        if row_no % 2 == 0:
            shade(cells[i], GRAY)
        margins(cells[i], 85, 95, 85, 95)
        p = cells[i].paragraphs[0]
        para(p, 0, 0, 1.08)
        font(p.add_run(value), 8.2, i == 0, BLUE if i == 0 else INK)
keep_rows_together(table)

heading(doc, "六、统一负面提示词")
text(doc, "写实摄影、3D塑料感、厚重油画、现代城市、科幻机械、枪械、完整魔法阵、完整圆月符号、暖金主光、血腥断肢、夸张伤口、性感姿势、暴露服装、哭泣求救、额外人物、幼崽、可见怪物、文字、水印、logo、UI边框、低清晰度、模糊脸、错误手指、多肢体、裁切头顶。", 9.5)

heading(doc, "七、音频生成提示词")
labeled(doc, "露娜音色基准：", "年轻女性，少女感但不幼态，声线清亮偏低，气息克制，警觉敏锐；负伤造成短促呼吸和轻微声带紧绷，但吐字仍准确。禁止哭腔、撒娇、虚弱呻吟、甜腻偶像声和英雄式高喊。")
audio_rows = [
    ("VO-01", "“不是野兽的味道……”", "近距离干声，露娜刚辨认完雾中气味，音量很轻，前半句是冷静判断，‘野兽’略微压低，尾音不做疑问上扬；句前有一次克制吸气，句后留1秒警戒静默。"),
    ("VO-02", "“断月纹还在追我。”", "近距离干声，肩伤疼痛但压住喘息，语速略慢；‘还在’表达确认威胁持续，‘追我’不是恐惧而是带怒意的事实判断。句尾收紧，不拖长。"),
    ("VO-03", "“有人把不属于森林的东西埋进来了。”", "中近距离干声，先看见古树刻痕再说话；‘有人’低沉警觉，‘不属于森林’强调边界被侵犯，最后一句压住愤怒并迅速恢复行动判断。"),
    ("AMB-01", "微光森林逆流雾环境", "90秒可循环立体声环境：黄昏森林风从左向右掠过树冠，贴地雾流却从右向左产生低频倒吸感；远处狼嚎两次，第二次更远；稀疏虫鸣，在月痕响应时瞬间静音。无鸟语乐园感，无恐怖尖叫。"),
    ("SFX-01", "脚步与负伤动作", "泥地轻脚步、踩断细枯枝、衣料摩擦、短促克制喘息；角色体型娇小，动作敏捷但落地因肩伤出现半步失衡。不要沉重盔甲声。"),
    ("SFX-02", "断月刻痕响应", "三段细小石裂声依次出现，随后进入银蓝能量低频共振；声音古老、冷、非机械，强度逐渐升高但不形成爆炸。"),
    ("SFX-03", "白光与村钟转场", "冷白能量在0.4秒内急速升高并骤停，硬切黑后远处村钟响第一声；钟体温暖但不喜庆，尾音自然延续4秒并进入下一场。"),
    ("BGM-01", "冷开场极简配乐", "约90秒，极简奇幻悬疑配乐，低音弦乐持续音、极少量玻璃质感泛音和不完整月主题动机；不出现完整旋律，不英雄化，不恐怖片化。60秒后逐渐增加张力，白光前达到峰值，切黑时完全停止，让村钟单独收尾。"),
]
table = doc.add_table(rows=1, cols=3)
geometry(table, [1100, 2300, 5960])
borders(table)
for i, value in enumerate(("编号", "素材", "生成提示词")):
    cell = table.rows[0].cells[i]
    shade(cell, BLUE)
    margins(cell)
    p = cell.paragraphs[0]
    para(p, 0, 0, 1.0, WD_ALIGN_PARAGRAPH.CENTER)
    font(p.add_run(value), 8.7, True, WHITE)
for row_no, values in enumerate(audio_rows, 1):
    cells = table.add_row().cells
    for i, value in enumerate(values):
        if row_no % 2 == 0:
            shade(cells[i], GRAY)
        margins(cells[i], 85, 95, 85, 95)
        p = cells[i].paragraphs[0]
        para(p, 0, 0, 1.08)
        font(p.add_run(value), 8.4, i == 0, BLUE if i == 0 else INK)
keep_rows_together(table)

video_heading = heading(doc, "八、视频生成提示词")
video_heading.paragraph_format.page_break_before = True
callout(doc, "生成策略", "每段建议生成5–10秒，再按逐镜头表剪辑。所有含露娜镜头必须固定使用同一角色参考图和相同随机种子/角色ID；背景镜头使用A01，刻痕与伤口分别用A03、A04作参考或后期叠加。")
video_rows = [
    ("V01", "S0-02", "A01背景", "横屏16:9，镜头从微光森林树冠缓慢下降到古树，树梢和落叶向画面右侧运动，贴地银蓝雾却向画面左侧逆流并汇入古树根部，黄昏冷青色，缓慢电影运镜，没有角色，没有怪物，保持场景结构稳定。"),
    ("V02", "S0-03/04", "A01+A02", "横屏16:9，使用露娜角色参考图保持脸、发型、狼耳、服装完全一致。低机位先拍靴子落地踉跄半步，再平滑抬升跟拍负伤露娜穿过林间；她左手压住肩伤，回头确认追踪者，随后看向古树，动作克制敏捷，发丝与斗篷受逆流雾影响。"),
    ("V03", "S0-05", "A01", "露娜主观视角缓慢扫过三处异常：无风自颤的叶片、突然中断的狼族足迹、向古树根部汇聚的雾线；依次拉焦，最后锁定树根。无角色正脸，无敌人实体，镜头稳定，不要快速晃动。"),
    ("V04", "S0-06/07", "A01+A02", "侧面中景转面部特写。露娜站到古树与森林深处之间，强迫身体站直，蹲身触地辨认气味后迅速收手；狼耳突然转向林中异响，金色眼睛收紧，护住伤肩但不后退。保留口型空间，禁止夸张表情和哭泣。"),
    ("V05", "S0-08", "A01+A03", "古树根部特写，贴地雾被树根缝隙吸入，三段冷白石刻线依次亮起并拼成残缺月牙，缺口清楚且像人为切断，伴随极少银蓝雾粒与石屑；缓慢推近，不出现完整法阵，不出现暖金光。"),
    ("V06", "S0-09", "A02+A03+A04", "露娜肩伤与古树刻痕的交叉特写。露娜伸手接近刻痕但尚未触碰，左肩伤口逸出断续银蓝雾光，被刻痕牵引成一条细线；她因疼痛短暂收紧手指但继续守住位置。特效低亮度、克制，不做技能爆发。"),
    ("V07", "S0-10", "A01+A02+A03+A06", "正面近景，残缺刻痕突然爆出冷白光，倒映在露娜金色瞳孔中；她迅速抬起受伤侧手臂护住面部，另一只手仍撑向古树前方保护身后领地。白光在0.4秒内吞没画面，随后硬切纯黑，禁止爆炸和镜头旋转。"),
]
table = doc.add_table(rows=1, cols=4)
geometry(table, [850, 1000, 1450, 6060])
borders(table)
for i, value in enumerate(("编号", "镜头", "参考素材", "视频生成提示词")):
    cell = table.rows[0].cells[i]
    shade(cell, BLUE)
    margins(cell)
    p = cell.paragraphs[0]
    para(p, 0, 0, 1.0, WD_ALIGN_PARAGRAPH.CENTER)
    font(p.add_run(value), 8.7, True, WHITE)
for row_no, values in enumerate(video_rows, 1):
    cells = table.add_row().cells
    for i, value in enumerate(values):
        if row_no % 2 == 0:
            shade(cells[i], GRAY)
        margins(cells[i], 85, 95, 85, 95)
        p = cells[i].paragraphs[0]
        para(p, 0, 0, 1.08)
        font(p.add_run(value), 8.2, i == 0, BLUE if i == 0 else INK)
keep_rows_together(table)
text(doc, "视频统一负面提示词：角色换脸、发色变化、耳朵消失或增生、服装变化、年龄变化、身高比例漂移、额外人物、可见怪物、镜头抖动、快速推拉、肢体穿模、错误手指、嘴型乱动、暖金主光、完整魔法阵、爆炸、文字、水印、logo、UI。", 9.5, False, MUTED, True)

heading(doc, "九、场次验收")
for item in [
    "总时长落在60–120秒，默认剪辑约90秒。",
    "清楚读出露娜负伤、雾流反常、古树刻痕异常和未知威胁未现身。",
    "三句V2.0关键对白、白光切黑、村钟和结尾字幕均保留。",
    "露娜始终体现判断与守护，不出现求救化或卖萌化表演。",
    "未提前泄露幕后者、源萝之力、月痕全貌、任务、奖励和契约。",
    "跳过后无状态丢失、重复结算、黑屏停滞或村钟重复。",
]:
    p = doc.add_paragraph(style="List Number")
    para(p, 0, 4, 1.18)
    font(p.add_run(item), 10)
callout(doc, "上线阻断", "如果无法清楚表现“树梢风向与贴地雾向相反”，或将负伤实体露娜误画成卡灵投影，本场叙事机制不成立。", warning=True)

doc.core_properties.title = "《斗萝大陆》分场脚本 S0 冷开场：月痕"
doc.core_properties.subject = "依据开局剧情设计 V2.0 的制作执行脚本"
doc.core_properties.author = "斗萝大陆项目组"
doc.save(OUT)
print(OUT)
