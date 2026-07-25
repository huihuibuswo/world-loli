from __future__ import annotations

from copy import deepcopy
from pathlib import Path
import shutil
import sys

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


def clear_runs(paragraph: Paragraph) -> None:
    for run in list(paragraph.runs):
        paragraph._element.remove(run._element)


def set_text(paragraph: Paragraph, text: str) -> None:
    rpr = deepcopy(paragraph.runs[0]._element.rPr) if paragraph.runs and paragraph.runs[0]._element.rPr is not None else None
    clear_runs(paragraph)
    run = paragraph.add_run(text)
    if rpr is not None:
        run._element.insert(0, rpr)


def set_labeled(paragraph: Paragraph, label: str, body: str) -> None:
    first_rpr = deepcopy(paragraph.runs[0]._element.rPr) if paragraph.runs and paragraph.runs[0]._element.rPr is not None else None
    second_rpr = deepcopy(paragraph.runs[1]._element.rPr) if len(paragraph.runs) > 1 and paragraph.runs[1]._element.rPr is not None else None
    clear_runs(paragraph)
    first = paragraph.add_run(label)
    if first_rpr is not None:
        first._element.insert(0, first_rpr)
    second = paragraph.add_run(body)
    if second_rpr is not None:
        second._element.insert(0, second_rpr)


def insert_after(anchor: Paragraph, text: str, style: str) -> Paragraph:
    element = OxmlElement("w:p")
    anchor._p.addnext(element)
    paragraph = Paragraph(element, anchor._parent)
    paragraph.style = style
    paragraph.add_run(text)
    return paragraph


def set_cell(table, row: int, col: int, text: str) -> None:
    paragraph = table.cell(row, col).paragraphs[0]
    set_text(paragraph, text)


def clone_row_format(source_row, target_row) -> None:
    if source_row._tr.trPr is not None:
        target_row._tr.insert(0, deepcopy(source_row._tr.trPr))
    for source_cell, target_cell in zip(source_row.cells, target_row.cells):
        if source_cell._tc.tcPr is not None:
            target_cell._tc.remove(target_cell._tc.tcPr)
            target_cell._tc.insert(0, deepcopy(source_cell._tc.tcPr))
        source_p = source_cell.paragraphs[0]
        target_p = target_cell.paragraphs[0]
        if source_p._p.pPr is not None:
            if target_p._p.pPr is not None:
                target_p._p.remove(target_p._p.pPr)
            target_p._p.insert(0, deepcopy(source_p._p.pPr))


def add_row(table, values: list[str]) -> None:
    source_row = table.rows[-1]
    row = table.add_row()
    clone_row_format(source_row, row)
    for index, value in enumerate(values):
        set_cell(table, len(table.rows) - 1, index, value)


def main() -> None:
    source = Path(sys.argv[1])
    output = Path(sys.argv[2])
    shutil.copy2(source, output)
    document = Document(output)
    paragraphs = document.paragraphs
    tables = document.tables

    set_text(paragraphs[1], "开局剧情设计 V2.0")
    set_labeled(
        paragraphs[3],
        "一句话体验  ",
        "玩家在晨曦村完成基础准备，离村后遇见已负伤并受失控月痕侵蚀的狼娘·露娜；误会战结束时露娜重伤濒临昏迷，将自身共鸣凝成完整卡灵托付给玩家。玩家带她回村疗伤，并在长期主线《月痕追迹》的首个阶段《逆流雾源》中完成调查、战斗与回报闭环。",
    )
    set_labeled(
        paragraphs[7],
        "核心判断  ",
        "开局仍从三件可玩的小事建立世界认知，但露娜不再只是碎片来源。她以受伤实体进入剧情，以共鸣卡灵进入战斗系统，并通过返村疗伤与《逆流雾源》把角色关系、新手奖励和长期主线连接起来。",
    )
    set_text(paragraphs[10], "卡牌代表关系留下的力量。训练教官赠出的《破绽识破》来自教学认可；露娜赠予的完整卡灵与《月牙撕裂》则来自她在重伤时主动建立的共鸣托付。")
    set_text(paragraphs[18], "黄昏将至，一名初到斗萝大陆的旅人沿村道抵达晨曦村。村口的引路灯在玩家经过时出现一次短促共鸣，灯芯里浮出陌生的月牙纹。村长没有立即解释，只请玩家先完成三项入村准备：带回一份晨曦暖茶、接受训练教官的实战检验、在村庄东缘完成一次采集记录。")
    set_text(paragraphs[20], "玩家离开晨曦村进入微光森林后，发现月光空地的植物突然闭合，一名肩侧带伤、行动明显失衡的银色狼耳少女挡在古树前。狼娘·露娜此前已被断月纹袭击，又因玩家身上的月牙共鸣误认玩家为污染源，双方只能通过卡牌战斗确认彼此力量的来源。")
    set_text(paragraphs[21], "露娜战败后旧伤与失控月痕同时恶化，几乎昏迷。主角停止攻击，以基础卡牌形成稳定回路。露娜确认主角并非污染者，把自身月痕凝成可陪伴战斗的完整卡灵投影，并托付长期主线《月痕追迹》。主角带实体露娜回晨曦村疗伤；安置完成后，玩家从疗养中的露娜处接取首个可玩阶段《逆流雾源》，返回森林调查三处证据、击败雾痕兽影并回村复命。")
    set_text(paragraphs[27], "让露娜的受伤、误判和战败赋予异常真实代价；她赠予卡灵是为了在实体无法行动时继续追查，而不是无条件倒贴。")
    set_labeled(
        paragraphs[32],
        "推荐强制顺序  ",
        "S1 → S3 → S2/S4 可交换 → S5 → S6 → S7。训练战必须早于露娜战；S7 在完成返村疗伤后开放，形成首个长期主线可玩闭环。",
    )
    set_text(paragraphs[39], "画面从黑场进入。微光森林深处，雾逆着树梢的风向向古树聚拢。银色狼耳少女扶住受伤的肩侧穿过林间，落地时踉跄半步，却仍挡在古树与狼族领地之间。")
    set_text(paragraphs[40], "露娜（压低喘息）：‘不是野兽的味道……断月纹还在追我。有人把不属于森林的东西埋进来了。’")
    set_text(paragraphs[41], "古树根部亮起残缺月牙刻痕。露娜伸手前，旧伤渗出银色雾光，刻痕随即放出白光。画面切黑，远处村钟响第一声。")
    set_labeled(
        paragraphs[66],
        "场景目标  ",
        "让玩家面对第一个带伤但仍有明确动机的可收集角色；通过误会战、战后救助和主动托付理解‘卡灵来自关系与选择’。  |  建议时长：12–18 分钟  |  触发：玩家离开晨曦村并到达月光空地",
    )
    set_text(paragraphs[67], "露娜从雾中落在玩家与古树之间，受伤的肩侧浮着断续银光。她闻到玩家身上的引路灯共鸣，看到牌匣边缘浮出的月牙纹，立即认定玩家与雾袭者有关。她的判断错误但有依据：相同纹路此前刚刚使她负伤，也正在伤害狼族领地。")
    set_text(paragraphs[68], "对话允许玩家指出伤势、说明调查目的或要求停手，但不能跳过首场关键战。选项改变露娜开场态度与首回合风格，全部收束到同一战斗、重伤救助和卡灵托付。")
    set_text(paragraphs[73], "露娜：‘咳……旧伤裂开了。月痕还在吞噬我的意识……’")
    set_text(paragraphs[74], "玩家：‘别再动了。我会用基础卡牌的共鸣先稳住它。’")
    set_text(paragraphs[75], "露娜：‘你的回路没有污染……是我认错了人。可我已经走不回安全的地方。’")
    set_text(paragraphs[76], "她把掌心按在月痕上，银蓝色共鸣凝成一枚完整卡灵印记；这只是露娜力量与意识的战斗投影，负伤的实体仍留在主角面前。")
    set_text(paragraphs[77], "露娜：‘收下它。它会代替现在的我与你并肩。污染源还在森林深处……替我追下去。这是长期委托——月痕追迹。’")
    set_text(paragraphs[78], "奖励反馈  固定新手奖励：完整‘狼娘·露娜’卡灵直接到账；《月牙撕裂》×2 自动加入当前启用套牌。该战不发放露娜碎片，重复结算只复用已拥有资产。")
    set_text(paragraphs[79], "S6. 返回晨曦村：疗伤与托付")
    set_labeled(
        paragraphs[80],
        "场景目标  ",
        "把重伤露娜安全带回晨曦村，完成疗养安置、序章结算和长期主线发布。  |  建议时长：5–7 分钟  |  触发：露娜战胜并完成卡灵托付",
    )
    set_text(paragraphs[81], "主角扶着即将昏迷的露娜返回晨曦村。村长与森林向导在村中疗养点接应，先处理旧伤和月痕侵蚀，再核对玩家带回的环境记录。露娜暂时无法以实体行动，但她赠予的共鸣卡灵可正常参战。")
    set_text(paragraphs[82], "露娜清醒片刻后发布长期主线《月痕追迹》，第一阶段为《逆流雾源》。序章完成界面只展示疗养收尾、完整露娜卡灵奖励和一个主线面板，随后主界面持续显示第一阶段目标。")
    set_text(paragraphs[88], "露娜的敏锐首先表现为带伤仍坚持守护领地；她的嘴硬来自责任与警惕，而不是无理攻击。战败后的赠予不是立即无条件信任，而是她在失去行动能力时，把继续追查的可能性托付给主角。后续通过疗养对白和《逆流雾源》逐步建立关系。")
    set_text(paragraphs[90], "玩家在序章中只有态度选择，不做路线分叉。选项改变少量回应、露娜首回合风格或隐藏信任标记，但全部收束到误会战、救助、卡灵托付与返村疗养；《逆流雾源》同样保持单一目标链，避免首版维护多套关卡和奖励。")
    set_labeled(
        paragraphs[95],
        "实现原则  ",
        "剧情只读取服务端已确认状态；客户端动画、镜头和对白可以重播，金币、卡牌、卡灵、证据、任务完成与战斗奖励不能重复结算。实体露娜的疗养状态和卡灵收藏状态必须分开表达。",
    )
    set_text(paragraphs[107], "结算页仍描述露娜碎片，或没有区分负伤实体露娜与共鸣卡灵投影。")
    set_text(paragraphs[109], "玩家无法在 85 分钟内完成序章与首个主线阶段，且耗时来自重复跑图或强制长对白。")
    set_text(paragraphs[117], "本设计综合项目现有卡灵、地图生态、角色表、卡牌战斗、NPC/剧情系统、完整 Demo 联调、晨曦村 NPC 功能、植物与 AI 对话设计，并以当前代码与已批准需求为准：新玩家初始套牌只含基础攻击与防御；露娜序章胜利直接授予完整卡灵与《月牙撕裂》×2；实体露娜回晨曦村疗养；《月痕追迹》至少包含可玩的第一阶段《逆流雾源》。")
    set_text(paragraphs[118], "V2.0 替代 V1.0 中‘露娜只给 3/30 碎片、认可不等于立即拥有’的旧结论。普通怪物与其他 Boss 仍沿用碎片合成规则；露娜仅作为序章剧情直招例外。后续第二阶段、新地图和幕后污染者揭晓继续延后。")

    anchor = paragraphs[82]
    additions = [
        ("S7. 月痕追迹·第一阶段：逆流雾源", "Heading 2"),
        ("场景目标  完成长期主线第一次可玩闭环，让露娜的托付立即变成行动。  |  建议时长：15–22 分钟  |  触发：晨曦村疗养安置完成", "Normal"),
        ("疗养中的露娜指出，袭击她的断月纹会让附近雾流出现短暂逆转。她请玩家先找森林向导确认第二处逆流位置，并强调不要追逐声音，只记录能被重复观察的证据。", "Normal"),
        ("接受《逆流雾源》后，森林向导在微光森林标记三处调查点：异常闭合的月光植物、突然中断的狼族足迹、附着断月纹的雾核。每处互动都由服务端幂等记录，刷新或切图不会丢失。", "Normal"),
        ("三处证据齐全后，雾核凝成剧情敌人‘雾痕兽影’。它不是普通狼族，而是污染模仿出的攻击性轮廓。玩家失败或中途退出按失败结算，但证据保留，可直接重新挑战。", "Normal"),
        ("首次胜利后，兽影崩解并留下无法自然形成的断月纹排列。玩家带记录返回晨曦村，与疗养中的露娜完成回报。", "Normal"),
        ("S7.1 可玩步骤", "Heading 3"),
        ("与疗养中的露娜交谈，接取《逆流雾源》。", "List Number"),
        ("与森林向导交谈，解锁微光森林三处固定证据点。", "List Number"),
        ("调查月光植物、狼族足迹和断月雾核，进度达到 3/3。", "List Number"),
        ("击败剧情敌人‘雾痕兽影’；失败和中途退出允许重试。", "List Number"),
        ("返回晨曦村向露娜回报，完成第一阶段。", "List Number"),
        ("S7.2 阶段收尾对白", "Heading 3"),
        ("露娜：‘这不是野兽留下的痕迹。有人在用断月纹模仿狼族的力量。’", "Normal"),
        ("玩家：‘雾痕兽影已经消散，但纹路指向森林更深处。’", "Normal"),
        ("露娜：‘先把记录交给向导。我需要一点时间恢复。等我能重新站起来，我们再追那个没有气味的人。’", "Normal"),
        ("系统：主线《月痕追迹》第一阶段完成。长期目标更新：追查操纵断月纹的人。", "Normal"),
    ]
    for text, style in additions:
        anchor = insert_after(anchor, text, style)

    # Metadata summary.
    set_cell(tables[0], 1, 1, "首登可玩的叙事化新手序章 + 长期主线第一阶段")
    set_cell(tables[0], 2, 1, "65–85 分钟；熟练玩家约 45–60 分钟")
    set_cell(tables[0], 5, 1, "完整露娜卡灵、《月牙撕裂》×2、返村疗养、《月痕追迹·逆流雾源》")
    set_cell(tables[0], 6, 1, "2026-07-25")

    set_cell(tables[1], 1, 1, "完成移动、交互、商店、任务、采集、战斗、卡灵奖励、疗养收尾和长期主线首阶段的闭环。")
    set_cell(tables[1], 2, 1, "新玩家初始卡组只含基础攻防；露娜首胜直接授予完整卡灵与《月牙撕裂》×2；服务端权威且幂等结算。")
    set_cell(tables[1], 3, 1, "不使用失忆、天选之子或百科旁白；露娜受伤、误判、赠予与疗养必须形成可理解因果。")
    set_cell(tables[1], 4, 1, "不实现第二个长期主线阶段、新地图、狼族聚落、幕后污染者揭晓、多结局或大型过场。")
    set_cell(tables[1], 5, 1, "玩家理解露娜为何攻击、为何赠予卡灵、实体为何回村疗养，并能完成《逆流雾源》的调查战斗闭环。")

    set_cell(tables[2], 6, 0, "托付")
    set_cell(tables[2], 6, 1, "她把无法继续的行动交给我")
    set_cell(tables[2], 6, 2, "重伤、共鸣卡灵、返村疗养和长期主线同时成立。")
    add_row(tables[2], ["追索", "我要验证她留下的线索", "三处证据、雾痕兽影与返村回报形成首阶段闭环。"])

    set_cell(tables[3], 4, 1, "玩家获得露娜的共鸣卡灵，但实体露娜重伤")
    set_cell(tables[3], 4, 2, "用战斗投影与疗养实体双线连接收藏系统和角色关系。")

    set_cell(tables[4], 6, 2, "误会战、救助、固定新手奖励")
    set_cell(tables[4], 6, 3, "露娜负伤真相、完整卡灵托付")
    set_cell(tables[4], 7, 1, "返回晨曦村：疗伤与托付")
    set_cell(tables[4], 7, 2, "返村、疗养安置、主线发布")
    set_cell(tables[4], 7, 3, "实体露娜留村，《月痕追迹》开启")
    set_cell(tables[4], 7, 4, "5–7 分")
    add_row(tables[4], ["S7", "逆流雾源", "接取、调查、剧情战、回报", "断月纹证据与长期主线首闭环", "15–22 分"])

    set_cell(tables[5], 4, 3, "战斗胜利并完成卡灵托付")
    set_cell(tables[5], 5, 0, "opening_return")
    set_cell(tables[5], 5, 1, "露娜卡灵已发放，实体露娜待送回村")
    set_cell(tables[5], 5, 2, "返回晨曦村、疗养安置")
    set_cell(tables[5], 5, 3, "完成序章并发布《逆流雾源》")
    set_cell(tables[5], 6, 0, "moon_trace_stage1")
    set_cell(tables[5], 6, 1, "序章完成并接取《逆流雾源》")
    set_cell(tables[5], 6, 2, "向导对话、3 处调查、雾痕兽影战、返村回报")
    set_cell(tables[5], 6, 3, "完成阶段一并保留长期目标")

    set_cell(tables[10], 3, 0, "‘你已经受伤了，别再被那道纹路控制。’")
    set_cell(tables[10], 3, 1, "‘别把观察当成关心……但你说得对，它正在变重。’")
    set_cell(tables[10], 3, 2, "展示伤势；战后救助对白更直接。")

    set_cell(tables[11], 1, 1, "‘别靠近！你身上的共鸣正在牵动我的伤口。’")
    set_cell(tables[11], 1, 2, "露娜带伤出战，界面显示断续月痕；不预览碎片奖励。")
    set_cell(tables[11], 3, 1, "旧伤渗出银色雾光；露娜短暂停手：‘这股共鸣……不是从你身上开始的。’")
    set_cell(tables[11], 4, 1, "露娜战败后无法站立，主角停止攻击并稳定月痕。")
    set_cell(tables[11], 4, 2, "服务端首胜直接发完整露娜卡灵与《月牙撕裂》×2；不发碎片。")

    set_cell(tables[12], 1, 0, "露娜")
    set_cell(tables[12], 1, 1, "‘卡灵投影会代替现在的我与你并肩。污染源还在森林深处……替我追下去。’")
    set_cell(tables[12], 2, 0, "玩家")
    set_cell(tables[12], 2, 1, "‘先别说了。我带你回晨曦村疗伤。’")
    set_cell(tables[12], 3, 0, "村长")
    set_cell(tables[12], 3, 1, "‘先救人，再谈月痕。向导，把东侧疗养间打开。’")
    set_cell(tables[12], 4, 0, "系统")
    set_cell(tables[12], 4, 1, "序章《雾中月痕》完成。主线《月痕追迹·逆流雾源》已开启。")

    set_cell(tables[13], 6, 1, "首个强角色钩子、误会战、卡灵托付与长期主线")
    set_cell(tables[13], 6, 2, "敏锐、带伤仍守护族群；托付后进入疗养与调查关系线。")
    set_cell(tables[13], 6, 3, "不要把赠予写成战败倒贴；不要让卡灵出现后实体露娜凭空消失。")

    set_cell(tables[14], 7, 0, "露娜卡灵")
    set_cell(tables[14], 7, 1, "重伤露娜主动凝聚共鸣投影")
    set_cell(tables[14], 7, 2, "完整卡灵与《月牙撕裂》×2 直接到账；幂等且不发碎片。")
    set_cell(tables[14], 8, 0, "长期主线")
    set_cell(tables[14], 8, 1, "露娜疗养期间托付玩家继续调查")
    set_cell(tables[14], 8, 2, "证据进度、战斗和回报均可恢复；阶段完成不关闭长期追踪。")

    set_cell(tables[15], 6, 2, "完整露娜卡灵 +《月牙撕裂》×2；记录负伤与返村状态")
    set_cell(tables[15], 6, 3, "重伤托付剧情、固定新手奖励")
    set_cell(tables[15], 7, 2, "完成序章，安置实体露娜，开启《逆流雾源》")
    set_cell(tables[15], 7, 3, "疗养收尾、主线卡片、自动存档")
    add_row(tables[15], ["moon_trace.stage1.accepted", "与疗养露娜交谈", "创建阶段一任务状态", "主线追踪显示向导目标"])
    add_row(tables[15], ["moon_trace.evidence.updated", "调查固定证据点", "幂等记录证据 ID，最高 3/3", "证据反馈与下一导航"])
    add_row(tables[15], ["moon_trace.shadow.completed", "首次击败雾痕兽影", "记录首胜并推进返村回报", "调查记录到账，不重复结算"])
    add_row(tables[15], ["moon_trace.stage1.completed", "向露娜提交调查记录", "完成阶段一并保留长期目标", "更新为‘追查操纵断月纹的人’"])

    set_cell(tables[16], 3, 0, "luna_contract_completed")
    set_cell(tables[16], 3, 2, "保证完整卡灵与签名卡奖励幂等。")
    set_cell(tables[16], 4, 0, "luna_recovery_state")
    set_cell(tables[16], 4, 1, "枚举")
    set_cell(tables[16], 4, 2, "区分返村途中、疗养中和可对话状态。")
    set_cell(tables[16], 5, 0, "moon_trace_stage")
    set_cell(tables[16], 5, 1, "枚举")
    set_cell(tables[16], 5, 2, "控制《逆流雾源》接取、调查、战斗、回报与完成。")
    add_row(tables[16], ["moon_trace_evidence", "集合", "记录三处证据事件 ID，防止刷新或重复互动丢失/叠加。"])
    add_row(tables[16], ["moon_trace_shadow_completed", "布尔", "保证雾痕兽影首胜结算幂等。"])

    set_cell(tables[17], 1, 0, "复用晨曦村 NPC、三项准备、商店、训练、采集、微光森林、战斗和任务追踪；实现露娜直招、返村疗养与《逆流雾源》完整闭环。")
    set_cell(tables[17], 1, 1, "完整源萝历史、狼族聚落、新地图、第二/第三条月痕关卡和幕后污染者揭晓。")
    set_cell(tables[17], 2, 0, "新增疗养露娜 NPC 状态、三处证据事件、雾痕兽影剧情敌人和长期主线阶段状态。")
    set_cell(tables[17], 2, 1, "配音、Live2D、抱持搬运动画、复杂同伴助战、分支地图和多结局。")

    set_cell(tables[18], 1, 1, "断线或重复请求导致露娜卡灵、签名卡或阶段一奖励重复发放")
    set_cell(tables[18], 1, 2, "资产与任务一致性破坏，可被主动利用")
    set_cell(tables[18], 1, 3, "战斗、契约、证据和阶段完成均使用唯一约束、锁和幂等状态检查。")
    set_cell(tables[18], 4, 1, "实体露娜与卡灵投影关系表达不清")
    set_cell(tables[18], 4, 2, "玩家误以为露娜被收进卡牌或出现两个独立角色")
    set_cell(tables[18], 4, 3, "战后与疗养对白固定说明卡灵是力量/意识投影，实体在村中疗伤。")
    add_row(tables[18], ["P1", "证据节点或雾痕兽影状态丢失", "主线软锁或重复调查", "证据 ID 集合持久化；失败保留证据；首胜后可直接返村。"])

    set_cell(tables[19], 2, 1, "移动、NPC、商店、任务、训练战、地图、采集、露娜战、直招奖励、疗养与《逆流雾源》全部串联。")
    set_cell(tables[19], 3, 1, "测试玩家能复述：露娜为何带伤攻击、为何赠予卡灵、实体为何回村，以及《逆流雾源》要调查什么。")
    set_cell(tables[19], 5, 0, "获取规则")
    set_cell(tables[19], 5, 1, "露娜首胜显示完整卡灵与《月牙撕裂》×2；无露娜碎片；实体疗养状态清晰。")
    set_cell(tables[19], 6, 1, "露娜战与雾痕兽影战失败、退出、刷新、重复提交均可恢复且不重复发奖。")
    set_cell(tables[19], 8, 1, "序章加首阶段中位完成时间 65–85 分钟；熟练玩家可在 60 分钟内完成。")
    add_row(tables[19], ["长期主线", "完成《逆流雾源》后主线仍显示下一目标，且本版本不要求第二阶段可玩。"])

    set_cell(tables[20], 5, 2, "‘你身上有那道月痕的味道。别再往前——它正在牵动我的伤口。’")
    set_cell(tables[20], 6, 1, "露娜")
    set_cell(tables[20], 6, 2, "‘咳……旧伤裂开了。收下这道月痕，它会化成我的卡灵投影。’")
    set_cell(tables[20], 7, 2, "‘先救人，再谈月痕。向导，把东侧疗养间打开。’")
    add_row(tables[20], ["阶段一接取", "露娜", "‘第二处逆流雾源还在。替我确认三件事：花、足迹，还有那枚雾核。’"])
    add_row(tables[20], ["阶段一收束", "露娜", "‘有人在用断月纹模仿狼族的力量。等我能站起来，我们再追下去。’"])

    document.core_properties.title = "《斗萝大陆》开局剧情设计 V2.0"
    document.core_properties.subject = "序章《雾中月痕》与长期主线第一阶段《逆流雾源》"
    document.core_properties.comments = "基于 V1.0 更新：露娜受伤直招、返村疗养、完整卡灵新手奖励与首个可玩长期主线阶段。"
    document.save(output)


if __name__ == "__main__":
    main()
