from __future__ import annotations

import json
import sys
import zipfile
from hashlib import sha256
from pathlib import Path

from docx import Document


def paragraph_data(paragraph, index: int) -> dict:
    return {
        "index": index,
        "style": paragraph.style.name if paragraph.style else None,
        "text": paragraph.text,
        "alignment": int(paragraph.alignment) if paragraph.alignment is not None else None,
        "runs": [
            {
                "text": run.text,
                "bold": run.bold,
                "italic": run.italic,
                "size_pt": run.font.size.pt if run.font.size else None,
                "font": run.font.name,
                "color": str(run.font.color.rgb) if run.font.color and run.font.color.rgb else None,
            }
            for run in paragraph.runs
        ],
    }


def main() -> None:
    source = Path(sys.argv[1])
    target = Path(sys.argv[2])
    document = Document(source)
    package = []
    with zipfile.ZipFile(source) as archive:
        for info in archive.infolist():
            data = archive.read(info.filename)
            package.append(
                {
                    "path": info.filename,
                    "size": len(data),
                    "sha256": sha256(data).hexdigest(),
                }
            )
    result = {
        "source": str(source.resolve()),
        "sha256": sha256(source.read_bytes()).hexdigest(),
        "paragraphs": [paragraph_data(p, i) for i, p in enumerate(document.paragraphs)],
        "tables": [
            {
                "index": ti,
                "rows": [[cell.text for cell in row.cells] for row in table.rows],
            }
            for ti, table in enumerate(document.tables)
        ],
        "headers": [
            [paragraph_data(p, i) for i, p in enumerate(section.header.paragraphs)]
            for section in document.sections
        ],
        "footers": [
            [paragraph_data(p, i) for i, p in enumerate(section.footer.paragraphs)]
            for section in document.sections
        ],
        "package": package,
    }
    target.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
    if len(sys.argv) > 3:
        lines = [f"Paragraphs: {len(result['paragraphs'])}", f"Tables: {len(result['tables'])}"]
        lines.extend(
            f"P{item['index']:03d} [{item['style']}] {item['text']}"
            for item in result["paragraphs"]
        )
        for table in result["tables"]:
            lines.append(f"TABLE {table['index']}")
            lines.extend(" | ".join(row) for row in table["rows"])
        Path(sys.argv[3]).write_text("\n".join(lines), encoding="utf-8")


if __name__ == "__main__":
    main()
